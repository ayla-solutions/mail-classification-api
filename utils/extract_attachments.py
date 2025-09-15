"""
- Fetch messages from Graph with full body (HTML + text)
- Robust attachment text extraction:
    * pdf: try text layer via pdfplumber; fallback to OCR (if poppler+tesseract available)
    * docx/xlsx/csv/html/txt/images: best-effort text extraction
- Output fields per message (id, subject, body_text/html, attachments, attachment_text)

NEW (2025-08-31):
- Delegated Graph (/me) support: no mailbox parameter; the token belongs to the signed-in user.
"""

# =========================
# Imports
# =========================
import os
import base64
from io import BytesIO
from typing import List, Tuple

import requests
import pdfplumber
from pdf2image import convert_from_bytes
from PIL import ImageEnhance, Image
import pytesseract
import pandas as pd
from docx import Document
from bs4 import BeautifulSoup
from datetime import datetime, timedelta, timezone

# =========================
# HTML → plain text
# =========================
def _html_to_text(html: str | None) -> str:
    if not html:
        return ""
    soup = BeautifulSoup(html, "lxml")
    return soup.get_text(separator=" ", strip=True)


# =========================
# Attachment extractors
# =========================
def _extract_pdf_text_layer(pdf_bytes: bytes) -> str:
    out = []
    try:
        with pdfplumber.open(BytesIO(pdf_bytes)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    out.append(t)
    except Exception as e:
        print(f"[pdfplumber failed] {e}")
    return "\n".join(out).strip()

def _ocr_pdf(pdf_bytes: bytes, dpi: int = 300) -> str:
    """
    OCR fallback for scanned PDFs (requires poppler & tesseract available in the container/host).
    If missing, this will catch exceptions and return an error marker string (non-fatal).
    """
    try:
        lines: List[str] = []
        images = convert_from_bytes(pdf_bytes, dpi=dpi)
        for img in images:
            img = img.convert("L")
            img = ImageEnhance.Contrast(img).enhance(2.0)
            lines.extend(pytesseract.image_to_string(img).split("\n"))
        return "\n".join(lines).strip()
    except Exception as e:
        return f"[ERROR OCR PDF: {e}]"

def _ocr_image(img_bytes: bytes) -> str:
    try:
        img = Image.open(BytesIO(img_bytes))
        img = img.convert("L")
        img = ImageEnhance.Contrast(img).enhance(2.0)
        return pytesseract.image_to_string(img)
    except Exception as e:
        return f"[ERROR OCR image: {e}]"

def _extract_docx(docx_bytes: bytes) -> str:
    try:
        doc = Document(BytesIO(docx_bytes))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(parts).strip()
    except Exception as e:
        return f"[ERROR reading DOCX: {e}]"

def _extract_xlsx(xlsx_bytes: bytes) -> str:
    try:
        out: List[str] = []
        xls = pd.ExcelFile(BytesIO(xlsx_bytes))
        for sheet in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet, dtype=str)
            out.append(f"=== Sheet: {sheet} ===")
            out.append(df.fillna("").to_string(index=False))
        return "\n".join(out).strip()
    except Exception as e:
        return f"[ERROR reading XLSX: {e}]"

def _extract_csv(csv_bytes: bytes) -> str:
    try:
        df = pd.read_csv(BytesIO(csv_bytes), dtype=str)
        return df.fillna("").to_string(index=False)
    except Exception as e:
        return f"[ERROR reading CSV: {e}]"

def _extract_html(html_bytes: bytes) -> str:
    try:
        return _html_to_text(html_bytes.decode(errors="ignore"))
    except Exception as e:
        return f"[ERROR reading HTML: {e}]"

# =========================
# Public: single-attachment extraction
# =========================
def extract_text_from_attachment(file_bytes: bytes, name: str) -> Tuple[str, str]:
    """
    Returns (extracted_text, method_tag)
    """
    n = (name or "").lower().strip()
    try:
        if n.endswith(".pdf"):
            text = _extract_pdf_text_layer(file_bytes)
            if text:
                return text, "pdf-text"
            return _ocr_pdf(file_bytes), "pdf-ocr"
        if n.endswith(".docx"):
            return _extract_docx(file_bytes), "docx"
        if n.endswith(".xlsx"):
            return _extract_xlsx(file_bytes), "xlsx"
        if n.endswith(".csv"):
            return _extract_csv(file_bytes), "csv"
        if n.endswith((".png", ".jpg", ".jpeg", ".tif", ".tiff", ".webp")):
            return _ocr_image(file_bytes), "image-ocr"
        if n.endswith((".htm", ".html")):
            return _extract_html(file_bytes), "html"
        if n.endswith(".txt"):
            try:
                return file_bytes.decode(errors="ignore"), "text"
            except Exception:
                return "[ERROR reading TXT]", "text"
        return "[Unsupported File Type]", "unknown"
    except Exception as e:
        return f"[ERROR reading attachment: {e}]", "unknown"


# =========================
# Graph helpers (delegated, /me)
# =========================
def _download_attachment_bytes(headers: dict, msg_id: str, att: dict) -> bytes:
    """
    Prefer 'contentBytes' from Graph; otherwise fetch the $value stream.
    """
    if "contentBytes" in att and att["contentBytes"]:
        return base64.b64decode(att["contentBytes"])
    url = f"https://graph.microsoft.com/v1.0/me/messages/{att['messageId'] if 'messageId' in att else msg_id}/attachments/{att['id']}/$value"
    resp = requests.get(url, headers=headers)
    resp.raise_for_status()
    return resp.content

def _get_full_message_body(headers: dict, msg_id: str) -> tuple[str, str]:
    """
    Returns (body_html, body_text) by explicitly selecting 'body'.
    """
    url = f"https://graph.microsoft.com/v1.0/me/messages/{msg_id}?$select=body,bodyPreview"
    try:
        r = requests.get(url, headers=headers)
        r.raise_for_status()
        body = (r.json() or {}).get("body", {})
        ctype = (body.get("contentType") or "").lower()
        content = body.get("content") or ""
        if ctype == "html":
            return content, _html_to_text(content)
        return "", content
    except Exception as e:
        print(f"[BODY FETCH WARN] {e}")
        return "", ""


# =========================
# Public: fetch messages with bodies & attachments (delegated /me)
# =========================
def fetch_messages_with_attachments(token: str, since_days: int = None) -> list[dict]:
    """
    Returns list of mail dicts with:
      id, sender, received_from, subject, received_at,
      body_preview, mail_body_html, mail_body_text,
      attachments (names), attachment_methods, attachment_text
    """

    headers = {"Authorization": f"Bearer {token}"}
    TOP_N = int(os.getenv("GRAPH_MAIL_TOP", "10"))

    base_url = "https://graph.microsoft.com/v1.0/me/messages"
    select = "$select=id,subject,from,receivedDateTime,bodyPreview"
    order  = "$orderby=receivedDateTime desc"

    # Apply date filter if since_days is set
    if since_days:
        since = (datetime.now(timezone.utc) - timedelta(days=since_days))
        since_str = since.strftime("%Y-%m-%dT%H:%M:%SZ")  # <-- FIX: Graph requires Zulu time
        url = f"{base_url}?$top={TOP_N}&{select}&{order}&$filter=receivedDateTime ge {since_str}"
    else:
        url = f"{base_url}?$top={TOP_N}&{select}&{order}"

    res = requests.get(url, headers=headers)
    res.raise_for_status()
    data = res.json()

    results = []
    for msg in data.get("value", []):
        msg_id = msg["id"]

        # (1) Full body
        body_html, body_text = _get_full_message_body(headers, msg_id)

        # (2) Attachments
        atts_url = f"https://graph.microsoft.com/v1.0/me/messages/{msg_id}/attachments"
        atts_resp = requests.get(atts_url, headers=headers)
        atts_resp.raise_for_status()
        atts = atts_resp.json().get("value", [])

        attachment_names: List[str] = []
        extracted_content: List[str] = []
        methods: List[str] = []

        for att in atts:
            # Only process file attachments (skip item/reference types)
            if att.get("@odata.type") != "#microsoft.graph.fileAttachment":
                attachment_names.append(att.get("name", ""))
                continue

            name = att.get("name", "")
            try:
                content = _download_attachment_bytes(headers, msg_id, att)
                text, method = extract_text_from_attachment(content, name)
            except Exception as e:
                text, method = f"[ERROR downloading/extracting {name}: {e}]", "unknown"

            attachment_names.append(name)
            methods.append(method)
            extracted_content.append(text)

        results.append({
            "id": msg.get("id"),
            "sender": msg.get("from", {}).get("emailAddress", {}).get("name", "Unknown"),
            "received_from": msg.get("from", {}).get("emailAddress", {}).get("address", "Unknown"),
            "subject": msg.get("subject"),
            "received_at": msg.get("receivedDateTime"),
            "body_preview": msg.get("bodyPreview") or "",
            "mail_body": msg.get("body", {}).get("content", ""),
            "attachments": attachment_names,
            "attachment_methods": methods,
            "attachment_text": "\n\n".join(extracted_content),
        })

    return results
